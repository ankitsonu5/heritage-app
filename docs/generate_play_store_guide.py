from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "generate-play-store-guide.ps1"
OUTPUT = ROOT / "docs" / "Heritage-Diagnostics-Play-Store-Guide.docx"
ICON = ROOT / "store-assets" / "play-store" / "app-icon-512.png"

document = Document()
section = document.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = document.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2"):
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor(116, 19, 33)

if ICON.exists():
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ICON), width=Inches(0.9))

title = document.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Heritage Diagnostics")
subtitle = document.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Google Play Store Publishing Guide")
meta = document.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Project-specific step-by-step checklist | Version 1.2 (Code 3) | Updated 1 August 2026")

started = False
pattern = re.compile(
    r"^\s*Add-(Heading|Text|Check)(?:\(|\s+)'(.+)'(?:,\s*(\d+))?\)?\s*$"
)
for line in SOURCE.read_text(encoding="utf-8").splitlines():
    if line.strip() == "try {":
        started = True
        continue
    if not started:
        continue
    match = pattern.match(line)
    if not match:
        continue
    kind, text, level = match.groups()
    text = text.replace('"', '"')
    if kind == "Heading":
        document.add_heading(text, level=int(level or 1))
    elif kind == "Check":
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run("[ ] " + text)
    else:
        document.add_paragraph(text)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Heritage Diagnostics - Play Store Submission Guide")

document.save(OUTPUT)
print(OUTPUT)
